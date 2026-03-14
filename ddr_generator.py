"""
DDR Report Generator — AI-Powered System
==========================================
Reads an Inspection PDF + Thermal PDF, extracts text and images,
sends them to Claude AI, and generates a structured Word DDR report.

Usage:
    python ddr_generator.py --inspection report.pdf --thermal thermal.pdf --output DDR_Output.docx

Author: AI Generalist Assignment Submission
"""

import os
import sys
import json
import base64
import argparse
import requests
import subprocess
import shutil
from pathlib import Path
from datetime import datetime
from io import BytesIO

# PDF & image handling
from pdf2image import convert_from_path
from PIL import Image

# Word document generation
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml

# ─────────────────────────────────────────────
#  CONFIGURATION
# ─────────────────────────────────────────────
ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"
MODEL = "claude-opus-4-5"
MAX_TOKENS = 8000
IMAGE_DPI = 120          # DPI for PDF → image conversion
MAX_PAGES_PER_PDF = 30   # Safety limit

# Colors (RGB)
COLOR_PRIMARY   = RGBColor(0x1F, 0x4E, 0x79)   # Dark blue
COLOR_ACCENT    = RGBColor(0xC5, 0x5A, 0x11)   # Orange
COLOR_HIGH      = RGBColor(0xCC, 0x33, 0x00)   # Red
COLOR_MEDIUM    = RGBColor(0xCC, 0x66, 0x00)   # Orange
COLOR_LOW       = RGBColor(0x00, 0xB0, 0x50)   # Green
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT_BG  = RGBColor(0xF2, 0xF2, 0xF2)
COLOR_SUB_BG    = RGBColor(0xD6, 0xE4, 0xF0)


# ─────────────────────────────────────────────
#  STEP 1 — PDF IMAGE EXTRACTOR
# ─────────────────────────────────────────────
class PDFExtractor:
    """Converts PDF pages to images for AI analysis."""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def extract_pages(self, pdf_path: str, prefix: str, dpi: int = IMAGE_DPI) -> list[str]:
        """
        Convert all pages of a PDF to JPEG images.
        Returns list of image file paths.
        """
        print(f"  📄 Extracting pages from: {Path(pdf_path).name}")
        img_paths = []

        # Use pdftoppm (fast, reliable)
        out_prefix = str(self.output_dir / prefix)
        cmd = ["pdftoppm", "-jpeg", f"-r", str(dpi), pdf_path, out_prefix]
        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            print(f"  ⚠️  pdftoppm warning: {result.stderr[:100]}")

        # Collect output images
        for f in sorted(self.output_dir.glob(f"{prefix}*.jpg")):
            img_paths.append(str(f))

        print(f"  ✅ Extracted {len(img_paths)} pages")
        return img_paths[:MAX_PAGES_PER_PDF]

    def image_to_base64(self, img_path: str, max_width: int = 1200) -> str:
        """Convert image file to base64 string, resizing if needed."""
        with Image.open(img_path) as img:
            # Resize if too large (saves API tokens)
            if img.width > max_width:
                ratio = max_width / img.width
                new_h = int(img.height * ratio)
                img = img.resize((max_width, new_h), Image.LANCZOS)

            buf = BytesIO()
            img.save(buf, format="JPEG", quality=75)
            return base64.standard_b64encode(buf.getvalue()).decode("utf-8")


# ─────────────────────────────────────────────
#  STEP 2 — AI ANALYZER (Claude API)
# ─────────────────────────────────────────────
class AIAnalyzer:
    """Sends PDF images to Claude and extracts structured DDR data."""

    def __init__(self, api_key: str):
        self.api_key = api_key
        self.headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }

    def _call_api(self, messages: list, system_prompt: str) -> str:
        """Make a call to Claude API and return text response."""
        payload = {
            "model": MODEL,
            "max_tokens": MAX_TOKENS,
            "system": system_prompt,
            "messages": messages,
        }
        response = requests.post(ANTHROPIC_API_URL, headers=self.headers, json=payload, timeout=120)
        if response.status_code != 200:
            raise Exception(f"API error {response.status_code}: {response.text[:300]}")
        return response.json()["content"][0]["text"]

    def analyze_inspection_report(self, page_images: list[str], extractor: PDFExtractor) -> dict:
        """
        Pass inspection report pages to Claude.
        Returns extracted structured data as dict.
        """
        print("  🤖 Sending inspection report to Claude AI...")

        # Build vision message with all page images
        content = [
            {
                "type": "text",
                "text": (
                    "You are analyzing a property inspection report. "
                    "Extract ALL available information carefully. "
                    "Return ONLY valid JSON — no markdown, no code blocks, no extra text.\n\n"
                    "Extract into this exact JSON structure:\n"
                    "{\n"
                    '  "property_info": {\n'
                    '    "property_type": "",\n'
                    '    "floors": "",\n'
                    '    "inspection_date": "",\n'
                    '    "inspected_by": "",\n'
                    '    "score": "",\n'
                    '    "flagged_items": "",\n'
                    '    "previous_audit": "",\n'
                    '    "previous_repair": "",\n'
                    '    "customer_name": "Not Available",\n'
                    '    "address": "Not Available"\n'
                    "  },\n"
                    '  "impacted_areas": [\n'
                    "    {\n"
                    '      "area_number": 1,\n'
                    '      "negative_side": "description of damage",\n'
                    '      "positive_side": "description of source/root area",\n'
                    '      "page_numbers_negative": [11],\n'
                    '      "page_numbers_positive": [12]\n'
                    "    }\n"
                    "  ],\n"
                    '  "checklist_findings": [\n'
                    '    {"item": "item name", "result": "Yes/No/Moderate/N/A"}\n'
                    "  ],\n"
                    '  "missing_info": ["list of any blank or unclear fields"]\n'
                    "}"
                ),
            }
        ]

        # Add page images (first 10 pages = cover + impacted areas)
        for i, img_path in enumerate(page_images[:10]):
            b64 = extractor.image_to_base64(img_path)
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": b64},
            })

        response_text = self._call_api(
            [{"role": "user", "content": content}],
            system_prompt=(
                "You are an expert property inspection analyst. "
                "Extract data accurately from inspection report images. "
                "If a field is blank or unclear, write 'Not Available'. "
                "Never invent data. Return only valid JSON."
            ),
        )

        # Clean and parse JSON
        response_text = response_text.strip()
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        response_text = response_text.strip()

        try:
            return json.loads(response_text)
        except json.JSONDecodeError as e:
            print(f"  ⚠️  JSON parse warning: {e}. Using fallback parser...")
            # Try to extract JSON from response
            start = response_text.find("{")
            end = response_text.rfind("}") + 1
            if start >= 0 and end > start:
                return json.loads(response_text[start:end])
            raise

    def analyze_thermal_report(self, page_images: list[str], extractor: PDFExtractor) -> dict:
        """
        Pass thermal report pages to Claude.
        Returns extracted thermal data as dict.
        """
        print("  🤖 Sending thermal report to Claude AI...")

        content = [
            {
                "type": "text",
                "text": (
                    "You are analyzing a thermal imaging report for a property inspection. "
                    "Extract ALL thermal readings from each page. "
                    "Return ONLY valid JSON — no markdown, no code blocks.\n\n"
                    "{\n"
                    '  "device": "",\n'
                    '  "serial_number": "",\n'
                    '  "inspection_date": "",\n'
                    '  "emissivity": "",\n'
                    '  "reflected_temperature": "",\n'
                    '  "thermal_readings": [\n'
                    "    {\n"
                    '      "image_id": "RB02380X.JPG",\n'
                    '      "page_number": 1,\n'
                    '      "hotspot_celsius": "28.8",\n'
                    '      "coldspot_celsius": "23.4",\n'
                    '      "delta_celsius": "5.4",\n'
                    '      "location_hint": "skirting/wall/ceiling area description from visible photo",\n'
                    '      "interpretation": "active moisture / no moisture / suspect area"\n'
                    "    }\n"
                    "  ]\n"
                    "}"
                ),
            }
        ]

        for i, img_path in enumerate(page_images[:10]):
            b64 = extractor.image_to_base64(img_path)
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": b64},
            })

        response_text = self._call_api(
            [{"role": "user", "content": content}],
            system_prompt=(
                "You are an expert thermal imaging analyst. "
                "Read each thermal image page carefully. "
                "Extract hotspot, coldspot temperatures and interpret moisture presence. "
                "Return only valid JSON."
            ),
        )

        response_text = response_text.strip()
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        response_text = response_text.strip()

        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            start = response_text.find("{")
            end = response_text.rfind("}") + 1
            if start >= 0 and end > start:
                return json.loads(response_text[start:end])
            raise

    def generate_ddr_content(self, inspection_data: dict, thermal_data: dict) -> dict:
        """
        Send extracted data to Claude and generate full DDR report content.
        Returns structured DDR content as dict.
        """
        print("  🤖 Generating DDR report content via Claude AI...")

        prompt = f"""
You are a senior property inspection engineer writing a Detailed Diagnostic Report (DDR) for a client.

Here is the extracted inspection data:
{json.dumps(inspection_data, indent=2)}

Here is the thermal imaging data:
{json.dumps(thermal_data, indent=2)}

Generate a complete, professional DDR report. Return ONLY valid JSON with this structure:

{{
  "property_issue_summary": {{
    "overview": "2-3 sentence overview of the property and inspection",
    "total_issues": 7,
    "key_findings": ["finding 1", "finding 2", "finding 3"]
  }},
  "area_wise_observations": [
    {{
      "area_number": 1,
      "area_name": "Hall",
      "negative_side": "description of damage visible",
      "positive_side": "description of root/source area",
      "thermal_reading": "Hotspot: X°C | Coldspot: Y°C | Delta: Z°C",
      "observation": "detailed 2-3 sentence observation combining visual + thermal",
      "inspection_page": 11,
      "thermal_page": 1
    }}
  ],
  "probable_root_causes": [
    {{
      "cause_number": 1,
      "title": "Root Cause Title",
      "description": "Detailed explanation"
    }}
  ],
  "severity_assessment": [
    {{
      "area": "area name",
      "issue": "brief issue description",
      "severity": "HIGH",
      "reasoning": "why this severity level"
    }}
  ],
  "recommended_actions": [
    {{
      "priority": "IMMEDIATE",
      "action": "action title",
      "detail": "detailed steps"
    }}
  ],
  "additional_notes": [
    "note 1", "note 2", "note 3"
  ],
  "missing_or_unclear_info": [
    {{
      "field": "field name",
      "status": "Not Available — reason"
    }}
  ]
}}

Rules:
- Use ONLY data from the provided documents. Do NOT invent facts.
- If something is missing, write "Not Available"
- Use simple, client-friendly language
- Severity must be HIGH, MEDIUM, or LOW
- Priority must be IMMEDIATE, SHORT_TERM, or LONG_TERM
"""

        response_text = self._call_api(
            [{"role": "user", "content": prompt}],
            system_prompt=(
                "You are an expert property inspection report writer. "
                "Generate accurate, professional DDR reports from inspection data. "
                "Never invent facts. Always write in simple, clear English. "
                "Return only valid JSON."
            ),
        )

        response_text = response_text.strip()
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        response_text = response_text.strip()

        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            start = response_text.find("{")
            end = response_text.rfind("}") + 1
            if start >= 0 and end > start:
                return json.loads(response_text[start:end])
            raise


# ─────────────────────────────────────────────
#  STEP 3 — WORD DOCUMENT BUILDER
# ─────────────────────────────────────────────
class DDRDocumentBuilder:
    """Builds the final Word .docx DDR report with images."""

    def __init__(self):
        self.doc = DocxDocument()
        self._setup_document()

    def _setup_document(self):
        """Set page margins and default styles."""
        from docx.shared import Cm
        section = self.doc.sections[0]
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Styling helpers ──────────────────────

    def _set_cell_bg(self, cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _add_section_heading(self, text: str, level: int = 1):
        """Add a styled section heading."""
        p = self.doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.bold = True

        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_WHITE
            run.font.name = "Arial"
            # Blue background via paragraph shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1F4E79")
            pPr.append(shd)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.left_indent  = Pt(8)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = "Arial"
            # Bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "C55A11")
            pBdr.append(bottom)
            pPr.append(pBdr)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(5)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_ACCENT
            run.font.name = "Arial"
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)

    def _add_para(self, text: str, bold: bool = False, italic: bool = False,
                  color: RGBColor = None, size: int = 11):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Arial"
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        return p

    def _add_bullet(self, text: str):
        p = self.doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Arial"
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    def _add_label_value(self, label: str, value: str):
        p = self.doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.color.rgb = COLOR_PRIMARY
        r1.font.size = Pt(11)
        r1.font.name = "Arial"
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.name = "Arial"
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    def _add_image(self, img_path: str, width_inches: float = 3.0, caption: str = ""):
        """Embed an image from file path."""
        if not img_path or not os.path.exists(img_path):
            self._add_para("[Image Not Available]", italic=True, color=RGBColor(0x88,0x88,0x88))
            return
        try:
            p = self.doc.add_paragraph()
            run = p.add_run()
            run.add_picture(img_path, width=Inches(width_inches))
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            if caption:
                cp = self.doc.add_paragraph()
                cr = cp.add_run(caption)
                cr.italic = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0x55,0x55,0x55)
                cp.paragraph_format.space_after = Pt(6)
        except Exception as e:
            self._add_para(f"[Image load error: {e}]", italic=True,
                           color=RGBColor(0x88,0x88,0x88))

    def _add_image_pair(self, img1: str, img2: str, cap1: str = "", cap2: str = ""):
        """Add two images side-by-side in a 2-col table."""
        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for i, (img_path, cap) in enumerate([(img1, cap1), (img2, cap2)]):
            cell = tbl.cell(0, i)
            cell.width = Inches(3.0)
            # Remove cell borders
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none")
                tcBorders.append(el)
            tcPr.append(tcBorders)

            p = cell.paragraphs[0]
            run = p.add_run()
            if img_path and os.path.exists(img_path):
                try:
                    run.add_picture(img_path, width=Inches(2.8))
                except Exception:
                    p.add_run("[Image Not Available]").italic = True
            else:
                p.add_run("[Image Not Available]").italic = True

            if cap:
                cp = cell.add_paragraph()
                cr = cp.add_run(cap)
                cr.italic = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0x55,0x55,0x55)

        self.doc.add_paragraph()

    def _add_divider(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "BFBFBF")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)

    # ── Section builders ─────────────────────

    def build_cover(self, property_info: dict, thermal_data: dict):
        """Build the cover/title section."""
        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("DETAILED DIAGNOSTIC REPORT")
        r.bold = True
        r.font.size = Pt(26)
        r.font.color.rgb = COLOR_PRIMARY
        r.font.name = "Arial"
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after  = Pt(4)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("DDR — Property Inspection with Thermal Analysis")
        r2.italic = True
        r2.font.size = Pt(13)
        r2.font.color.rgb = COLOR_ACCENT
        r2.font.name = "Arial"
        p2.paragraph_format.space_after = Pt(16)

        # Info table
        info_rows = [
            ("Property Type",        property_info.get("property_type", "Not Available")),
            ("Building Floors",       property_info.get("floors", "Not Available")),
            ("Inspection Date & Time",property_info.get("inspection_date", "Not Available")),
            ("Inspected By",          property_info.get("inspected_by", "Not Available")),
            ("Thermal Camera Used",   f"{thermal_data.get('device','Not Available')} (SN: {thermal_data.get('serial_number','N/A')})"),
            ("Inspection Score",      property_info.get("score", "Not Available")),
            ("Previous Structural Audit", property_info.get("previous_audit", "Not Available")),
            ("Previous Repair Work",  property_info.get("previous_repair", "Not Available")),
        ]

        tbl = self.doc.add_table(rows=len(info_rows), cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (label, value) in enumerate(info_rows):
            row = tbl.rows[i]
            # Label cell
            lc = row.cells[0]
            lc.width = Inches(3.0)
            lp = lc.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(11)
            lr.font.color.rgb = COLOR_PRIMARY
            lr.font.name = "Arial"
            self._set_cell_bg(lc, "D6E4F0")

            # Value cell
            vc = row.cells[1]
            vc.width = Inches(3.5)
            vp = vc.paragraphs[0]
            vr = vp.add_run(value)
            vr.font.size = Pt(11)
            vr.font.name = "Arial"
            if "No" in value or "Not Available" in value:
                vr.font.color.rgb = COLOR_HIGH

        self.doc.add_paragraph()
        self.doc.add_page_break()

    def build_section1(self, summary_data: dict, areas: list):
        """Section 1: Property Issue Summary"""
        self._add_section_heading("1. Property Issue Summary")
        self._add_para(summary_data.get("overview", ""))

        # Key findings
        findings = summary_data.get("key_findings", [])
        if findings:
            self._add_para("Key Findings:", bold=True)
            for f in findings:
                self._add_bullet(f)

        self._add_section_heading("Summary of Issues", level=3)

        # Summary table
        tbl = self.doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"

        # Header row
        headers = ["No.", "Impacted Area (–ve Side)", "Root/Source Area (+ve Side)", "Severity"]
        widths  = [0.4, 2.7, 2.7, 0.9]
        for i, (h, w) in enumerate(zip(headers, widths)):
            cell = tbl.cell(0, i)
            cell.width = Inches(w)
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.color.rgb = COLOR_WHITE
            r.font.size = Pt(10)
            self._set_cell_bg(cell, "1F4E79")

        # Data rows
        sev_colors = {"HIGH": "CC3300", "MEDIUM": "CC6600", "LOW": "00B050"}
        for area in areas:
            row = tbl.add_row()
            cells = row.cells
            data = [
                str(area.get("area_number", "")),
                area.get("negative_side", ""),
                area.get("positive_side", ""),
                area.get("severity", "HIGH"),
            ]
            for i, (val, w) in enumerate(zip(data, widths)):
                c = cells[i]
                c.width = Inches(w)
                p = c.paragraphs[0]
                r = p.add_run(val)
                r.font.size = Pt(10)
                if i == 3:
                    r.bold = True
                    r.font.color.rgb = COLOR_WHITE
                    sev = val.upper()
                    self._set_cell_bg(c, sev_colors.get(sev, "888888"))

        self.doc.add_paragraph()
        self._add_divider()

    def build_section2(self, areas: list, insp_images: list, thermal_images: list):
        """Section 2: Area-Wise Observations with images"""
        self._add_section_heading("2. Area-Wise Observations")
        self._add_para(
            "Each impacted area is described with visual observations and thermal "
            "imaging findings. Site photographs and thermal images are included."
        )

        for area in areas:
            num  = area.get("area_number", "")
            name = area.get("area_name", f"Area {num}")
            self._add_section_heading(f"Area {num} — {name}", level=2)

            self._add_label_value("Negative Side", area.get("negative_side", "Not Available"))
            self._add_label_value("Positive Side",  area.get("positive_side", "Not Available"))
            self._add_label_value("Thermal Reading",area.get("thermal_reading", "Not Available"))
            self._add_para(area.get("observation", ""))

            # Site photographs
            self._add_section_heading("Site Photographs", level=3)
            insp_pg = area.get("inspection_page", None)
            if insp_pg and insp_pg - 1 < len(insp_images):
                img1 = insp_images[insp_pg - 1] if insp_pg - 1 >= 0 else None
                img2 = insp_images[insp_pg] if insp_pg < len(insp_images) else None
                self._add_image_pair(
                    img1, img2,
                    f"Inspection — {name} (negative side)",
                    f"Inspection — {name} (source area)"
                )
            else:
                self._add_para("[Image Not Available]", italic=True,
                               color=RGBColor(0x88,0x88,0x88))

            # Thermal image
            self._add_section_heading("Thermal Image", level=3)
            thermal_pg = area.get("thermal_page", None)
            if thermal_pg and thermal_pg - 1 < len(thermal_images):
                t_img = thermal_images[thermal_pg - 1]
                self._add_image(t_img, width_inches=3.5,
                                caption=f"Thermal scan — {area.get('thermal_reading','')}")
            else:
                self._add_para("[Thermal Image Not Available]", italic=True,
                               color=RGBColor(0x88,0x88,0x88))

            self._add_divider()

    def build_section3(self, root_causes: list):
        """Section 3: Probable Root Cause"""
        self._add_section_heading("3. Probable Root Cause Analysis")
        self._add_para(
            "Based on combined analysis of visual inspection data and thermal imaging, "
            "the following root causes have been identified:"
        )
        for rc in root_causes:
            self._add_section_heading(
                f"Cause {rc.get('cause_number','')} — {rc.get('title','')}", level=3
            )
            self._add_para(rc.get("description", ""))
        self._add_divider()

    def build_section4(self, severity_data: list):
        """Section 4: Severity Assessment"""
        self._add_section_heading("4. Severity Assessment")
        self._add_para(
            "Severity levels are assigned based on extent of damage, structural risk, "
            "health concerns and urgency of intervention."
        )

        tbl = self.doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        headers = ["Area / Issue", "Observation", "Severity", "Reasoning"]
        widths  = [1.8, 2.5, 0.9, 2.0]
        for i, (h, w) in enumerate(zip(headers, widths)):
            c = tbl.cell(0, i)
            c.width = Inches(w)
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.color.rgb = COLOR_WHITE
            r.font.size = Pt(10)
            self._set_cell_bg(c, "1F4E79")

        sev_colors = {"HIGH": "CC3300", "MEDIUM": "CC6600", "LOW": "00B050"}
        for i, item in enumerate(severity_data):
            row = tbl.add_row()
            bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
            vals = [
                item.get("area", ""),
                item.get("issue", ""),
                item.get("severity", "HIGH"),
                item.get("reasoning", ""),
            ]
            for j, (val, w) in enumerate(zip(vals, widths)):
                c = row.cells[j]
                c.width = Inches(w)
                r = c.paragraphs[0].add_run(val)
                r.font.size = Pt(10)
                if j == 2:
                    r.bold = True
                    r.font.color.rgb = COLOR_WHITE
                    self._set_cell_bg(c, sev_colors.get(val.upper(), "888888"))
                else:
                    self._set_cell_bg(c, bg)

        self.doc.add_paragraph()
        self._add_divider()

    def build_section5(self, actions: list):
        """Section 5: Recommended Actions"""
        self._add_section_heading("5. Recommended Actions")

        priority_order = {"IMMEDIATE": 1, "SHORT_TERM": 2, "LONG_TERM": 3}
        priority_labels = {
            "IMMEDIATE":   ("Immediate Priority (Within 1–2 Weeks)", "CC3300"),
            "SHORT_TERM":  ("Short-Term Priority (Within 1 Month)",  "CC6600"),
            "LONG_TERM":   ("Long-Term Recommendation",              "1F4E79"),
        }

        sorted_actions = sorted(actions, key=lambda x: priority_order.get(x.get("priority","LONG_TERM"), 3))
        current_priority = None

        for action in sorted_actions:
            p = action.get("priority", "LONG_TERM")
            if p != current_priority:
                label, color = priority_labels.get(p, ("Other", "888888"))
                self._add_section_heading(label, level=3)
                current_priority = p
            self._add_bullet(f"{action.get('action','')}: {action.get('detail','')}")

        self._add_divider()

    def build_section6(self, notes: list):
        """Section 6: Additional Notes"""
        self._add_section_heading("6. Additional Notes")
        for note in notes:
            self._add_bullet(note)
        self._add_divider()

    def build_section7(self, missing_info: list):
        """Section 7: Missing or Unclear Information"""
        self._add_section_heading("7. Missing or Unclear Information")
        self._add_para(
            "The following information was not available in the provided documents "
            "or could not be confirmed from the inspection data:"
        )

        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        for i, h in enumerate(["Information Type", "Status"]):
            c = tbl.cell(0, i)
            c.width = Inches(2.5 if i == 0 else 4.7)
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.color.rgb = COLOR_WHITE
            r.font.size = Pt(10)
            self._set_cell_bg(c, "1F4E79")

        for i, item in enumerate(missing_info):
            row = tbl.add_row()
            bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
            field  = item.get("field", "")
            status = item.get("status", "Not Available")

            c0 = row.cells[0]
            c0.width = Inches(2.5)
            r0 = c0.paragraphs[0].add_run(field)
            r0.bold = True
            r0.font.size = Pt(10)
            self._set_cell_bg(c0, bg)

            c1 = row.cells[1]
            c1.width = Inches(4.7)
            r1 = c1.paragraphs[0].add_run(status)
            r1.font.size = Pt(10)
            if "Not Available" in status:
                r1.font.color.rgb = COLOR_HIGH
            self._set_cell_bg(c1, bg)

        # Footer note
        self.doc.add_paragraph()
        fn = self.doc.add_paragraph()
        fnr = fn.add_run(
            "This report was generated by an AI-powered DDR system using Claude AI. "
            "All findings are based solely on the provided inspection and thermal documents. "
            "This report should be reviewed by a qualified structural/waterproofing engineer "
            "before initiating any repair works."
        )
        fnr.italic = True
        fnr.font.size = Pt(9)
        fnr.font.color.rgb = RGBColor(0x88,0x88,0x88)

    def save(self, output_path: str):
        """Save the document to file."""
        self.doc.save(output_path)
        print(f"  ✅ Document saved: {output_path}")


# ─────────────────────────────────────────────
#  MAIN PIPELINE
# ─────────────────────────────────────────────
def run_pipeline(
    inspection_pdf: str,
    thermal_pdf: str,
    output_path: str,
    api_key: str,
    work_dir: str = "./extracted_images",
):
    """
    Full DDR generation pipeline:
    1. Extract images from PDFs
    2. Send to Claude AI for analysis
    3. Generate DDR Word document
    """
    print("\n" + "="*60)
    print("  DDR Report Generator — AI Pipeline")
    print("="*60)

    # ── Step 1: Extract images ─────────────────
    print("\n[STEP 1] Extracting images from PDFs...")
    extractor = PDFExtractor(work_dir)
    insp_images    = extractor.extract_pages(inspection_pdf, "insp_page")
    thermal_images = extractor.extract_pages(thermal_pdf,    "thermal_page")

    # ── Step 2: AI Analysis ────────────────────
    print("\n[STEP 2] Analyzing documents with Claude AI...")
    analyzer = AIAnalyzer(api_key)

    inspection_data = analyzer.analyze_inspection_report(insp_images, extractor)
    thermal_data    = analyzer.analyze_thermal_report(thermal_images, extractor)

    # Save extracted data for debugging/audit
    with open(os.path.join(work_dir, "inspection_data.json"), "w") as f:
        json.dump(inspection_data, f, indent=2)
    with open(os.path.join(work_dir, "thermal_data.json"), "w") as f:
        json.dump(thermal_data, f, indent=2)
    print("  ✅ Raw extracted data saved to extracted_images/*.json")

    # Merge thermal readings into area data
    thermal_readings = thermal_data.get("thermal_readings", [])
    areas = inspection_data.get("impacted_areas", [])
    for i, area in enumerate(areas):
        if i < len(thermal_readings):
            tr = thermal_readings[i]
            hot  = tr.get("hotspot_celsius", "N/A")
            cold = tr.get("coldspot_celsius", "N/A")
            delta = tr.get("delta_celsius", "N/A")
            area["thermal_reading"] = f"Hotspot: {hot}°C | Coldspot: {cold}°C | Delta: {delta}°C"
            area["thermal_page"] = tr.get("page_number", i + 1)

    # Generate full DDR content via AI
    ddr_content = analyzer.generate_ddr_content(inspection_data, thermal_data)

    # Merge area severity into area list
    severity_map = {
        s.get("area", "").lower(): s.get("severity", "HIGH")
        for s in ddr_content.get("severity_assessment", [])
    }
    for area in ddr_content.get("area_wise_observations", areas):
        area_name = area.get("area_name", "").lower()
        area["severity"] = severity_map.get(area_name, "HIGH")
        if "thermal_reading" not in area:
            for a in areas:
                if a.get("area_number") == area.get("area_number"):
                    area["thermal_reading"] = a.get("thermal_reading", "Not Available")
                    area["thermal_page"]    = a.get("thermal_page", 1)

    # ── Step 3: Build Word Document ────────────
    print("\n[STEP 3] Building Word document...")
    builder = DDRDocumentBuilder()

    prop_info = inspection_data.get("property_info", {})
    ddr_areas = ddr_content.get("area_wise_observations", areas)

    builder.build_cover(prop_info, thermal_data)
    builder.build_section1(ddr_content.get("property_issue_summary", {}), ddr_areas)
    builder.build_section2(ddr_areas, insp_images, thermal_images)
    builder.build_section3(ddr_content.get("probable_root_causes", []))
    builder.build_section4(ddr_content.get("severity_assessment", []))
    builder.build_section5(ddr_content.get("recommended_actions", []))
    builder.build_section6(ddr_content.get("additional_notes", []))
    builder.build_section7(ddr_content.get("missing_or_unclear_info", []))
    builder.save(output_path)

    print("\n" + "="*60)
    print(f"  ✅ DDR Report generated successfully!")
    print(f"  📄 Output: {output_path}")
    print("="*60 + "\n")
    return output_path


# ─────────────────────────────────────────────
#  CLI ENTRY POINT
# ─────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="AI-powered DDR Report Generator",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python ddr_generator.py --inspection Sample_Report.pdf --thermal Thermal_Images.pdf
  python ddr_generator.py --inspection report.pdf --thermal thermal.pdf --output MyReport.docx
  python ddr_generator.py --inspection report.pdf --thermal thermal.pdf --api-key sk-ant-...
        """,
    )
    parser.add_argument("--inspection", required=True, help="Path to inspection report PDF")
    parser.add_argument("--thermal",    required=True, help="Path to thermal images PDF")
    parser.add_argument("--output",     default="DDR_Report_Output.docx", help="Output .docx path")
    parser.add_argument("--api-key",    default=None, help="Anthropic API key (or set ANTHROPIC_API_KEY env var)")
    parser.add_argument("--work-dir",   default="./extracted_images", help="Working directory for temp images")

    args = parser.parse_args()

    # Get API key
    api_key = args.api_key or os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: Anthropic API key required.")
        print("  Set ANTHROPIC_API_KEY environment variable, or use --api-key flag.")
        sys.exit(1)

    # Validate inputs
    if not os.path.exists(args.inspection):
        print(f"ERROR: Inspection PDF not found: {args.inspection}")
        sys.exit(1)
    if not os.path.exists(args.thermal):
        print(f"ERROR: Thermal PDF not found: {args.thermal}")
        sys.exit(1)

    run_pipeline(
        inspection_pdf=args.inspection,
        thermal_pdf=args.thermal,
        output_path=args.output,
        api_key=api_key,
        work_dir=args.work_dir,
    )
